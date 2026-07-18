from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path("fixtures/uploads")

def create(filename, student, course, title, sections):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    for name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor(46, 116, 181)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{student}\n{course}\n18 July 2026").italic = True
    for section_title, paragraphs in sections:
        doc.add_heading(section_title, level=1)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
    ROOT.mkdir(parents=True, exist_ok=True)
    doc.save(ROOT / filename)

create("hamza-ali-ai-feedback-essay.docx", "Hamza Ali", "English Composition", "Should Students Use AI Feedback on Early Drafts?", [
    ("Introduction", ["Students should be allowed to use AI feedback on an early essay draft when the use is disclosed and the student remains responsible for every revision. A feedback tool can help a writer notice unclear sentences, missing transitions, and repeated words, but it cannot replace the judgment needed to decide what an argument means or whether evidence is convincing.", "The strongest reason to permit limited AI feedback is that students already use spell-checkers, writing centres, and peer comments. These supports point out possible problems without submitting the final work for the student. A clear classroom rule can distinguish feedback from generating a finished answer." ]),
    ("Benefits and limits", ["For a student who is still learning academic English, immediate feedback can make revision less intimidating. The student can compare several suggestions, reject weak ones, and explain which changes improved the draft. This process can strengthen revision skills when teachers ask for a short reflection on the final choices.", "However, AI feedback can sound confident even when it misunderstands a source or suggests an overly formal voice. Students should not copy a suggested sentence without checking it against their own meaning. Teachers should also continue to teach citation, paraphrasing, and the difference between a claim and evidence." ]),
    ("A practical policy", ["Schools could allow AI feedback before final submission if students attach a brief record of the prompts they used and describe two changes they accepted or rejected. The teacher would then see the student's reasoning rather than assuming that every polished sentence came from the same source.", "This policy would be fairer than either banning every tool or allowing invisible use. It makes the student accountable for the final essay while recognizing that thoughtful feedback can be part of a normal drafting process." ]),
    ("Conclusion", ["AI feedback should be treated as a limited revision aid, not as an author. Disclosure and reflection allow teachers to assess the student's own decisions while giving students a practical way to improve early drafts."])
])

create("sana-rahman-school-gardens-essay.docx", "Sana Rahman", "Environmental Studies", "Why Schools Should Create Community Gardens", [
    ("Introduction", ["Schools should create small community gardens because they give students a practical way to learn science, improve unused spaces, and support healthier eating habits. A garden is not a solution to every environmental problem, but it can make lessons about soil, water, food, and local biodiversity more concrete.", "Students often learn about ecosystems from diagrams alone. When they measure plant growth, observe insects, and plan watering schedules, they can connect those ideas to a real shared responsibility." ]),
    ("Learning and wellbeing", ["A school garden can support science lessons by giving classes a place to test how sunlight, compost, and water affect plants. It can also support mathematics when students calculate planting areas, track harvest weights, or compare the cost of seeds with the value of produce.", "Gardening may also give students a quieter outdoor activity during a busy school day. This benefit should not be exaggerated, but regular work in a shared green space can help students feel ownership of their school environment." ]),
    ("Costs and access", ["Critics may argue that gardens require money, staff time, and water. Those concerns are valid because a neglected garden can become another burden for teachers. Schools should begin with raised beds or containers, use drought-tolerant plants, and involve families or local organisations in a realistic maintenance plan.", "Access also matters. A garden should include paths and work areas that students with different mobility needs can use. Planning for access at the start makes the project more inclusive and prevents the garden from benefiting only a small group." ]),
    ("Conclusion", ["A modest, well-planned school garden can make learning more active and improve the school environment. Starting small, sharing responsibility, and designing for access would make the project more likely to succeed."])
])
