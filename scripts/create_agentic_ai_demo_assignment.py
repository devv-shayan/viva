from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("fixtures/uploads/sana-rahman-agentic-ai-demo-essay.docx")


def set_font(run, name="Calibri", size=11, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    set_font(paragraph.add_run(text))


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(title.add_run("Should Universities Use Agentic AI Study Assistants?"), size=20, bold=True, color=(11, 37, 69))

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(18)
    set_font(metadata.add_run("Sana Rahman | Introduction to Computing | Demo assignment"), size=10, color=(89, 89, 89))

    add_body(doc, "Universities should allow a limited agentic AI study assistant in introductory courses. Agentic AI can plan small steps, find sources, and ask students to check their own reasoning. It should support learning, not complete a student’s final work.")

    add_body(doc, "First, the assistant can make difficult work less overwhelming. For example, a student who is stuck on a programming task could receive a three-step plan: identify the error, read the relevant course note, and test one fix. The assistant should show where its information came from instead of giving an unexplained answer.")

    add_body(doc, "However, there are real risks. Students may rely on the tool too much, accept an incorrect answer, or share private course data. Universities should set boundaries: the assistant must not write final submissions, it must show sources and uncertainty, and teachers should be able to review an activity record when needed.")

    add_body(doc, "Some people argue that universities should ban agentic AI because rules are easier to enforce that way. I disagree. A complete ban may push use outside the classroom, where teachers cannot guide it. A transparent, limited pilot is fairer because students can learn responsible use and teachers can see what support was given.")

    add_body(doc, "In conclusion, a university should test agentic AI with a small pilot. It should measure whether students participate more, make fewer repeated mistakes, and still explain their own work in a short viva. Teachers should keep the final decision about grades and feedback.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Demo assignment - fictional content"), size=9, color=(120, 120, 120))

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
