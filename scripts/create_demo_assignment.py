from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

out = Path("fixtures/uploads/areeba-khan-congestion-pricing-essay.docx")
out.parent.mkdir(parents=True, exist_ok=True)
doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15
for name, size in [("Heading 1", 16), ("Heading 2", 13)]:
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(46, 116, 181)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Should Karachi Adopt Congestion Pricing?")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(23, 23, 23)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Areeba Khan\nCivics and Public Policy\n18 July 2026").italic = True

doc.add_heading("Introduction", level=1)
for paragraph in [
    "Karachi faces long travel times, unreliable buses, and worsening air quality as more cars enter the city centre each year. Congestion pricing, which charges drivers for entering the busiest areas during peak hours, could reduce unnecessary traffic and create a stable source of funding for public transport. Karachi should adopt a carefully designed congestion charge because evidence from other cities shows that it can reduce traffic, while exemptions and transit investment can protect lower-income commuters.",
    "London introduced its congestion charge in 2003 and saw traffic in the charging zone fall by around 15 percent, with journey-time reliability improving in the first years of the scheme. Stockholm also used a trial before making its charge permanent; public acceptance increased after residents experienced shorter travel times and clearer information about how the revenue would be spent. These examples do not prove that Karachi would have identical results, but they show that pricing can change travel choices when people have realistic alternatives.",
]: doc.add_paragraph(paragraph)

doc.add_heading("Equity and public transport", level=1)
for paragraph in [
    "Critics argue that congestion pricing falls hardest on low-income commuters, who may have no alternative to driving. This objection matters because a charge would be unfair if it simply added a cost without improving mobility. Karachi should therefore provide exemptions or discounts for people with disabilities, essential workers on late shifts, and households that can demonstrate limited access to public transport.",
    "Revenue should be publicly reported and directed toward more frequent buses, safer walking routes, and reliable connections to employment areas. If the policy helps buses move faster and makes them more dependable, it can benefit people who already use public transport as well as people who choose to leave their cars at home. The policy should be judged not only by traffic reduction, but also by whether these improvements are actually delivered.",
]: doc.add_paragraph(paragraph)

doc.add_heading("A practical pilot", level=1)
for paragraph in [
    "Karachi could begin with a two-year pilot in the central business district rather than applying a citywide charge immediately. Existing Safe City ANPR infrastructure could make enforcement more feasible, although the city would need to publish clear privacy rules, an appeal process, and an estimate of operating costs. A pilot would allow officials to measure traffic, bus speeds, revenue, and the effect on different neighbourhoods before deciding whether to expand the programme.",
    "A charge of 200 rupees during the busiest hours could generate substantial revenue, but the exact amount depends on the number of vehicles, exemptions, and compliance rates. For that reason, the city should not promise a specific total before publishing a transparent calculation. A limited pilot with regular independent reporting would be a more responsible way to test the policy than making broad claims without local evidence.",
]: doc.add_paragraph(paragraph)

doc.add_heading("Conclusion", level=1)
doc.add_paragraph("Karachi should test congestion pricing through a transparent pilot with clear exemptions and a binding commitment to improve public transport. The strongest case for the policy is not that it punishes drivers, but that it can make travel more reliable and finance alternatives to driving. If the evidence from the pilot does not show those benefits, the city should revise or stop the programme.")

doc.save(out)
print(out.resolve())
